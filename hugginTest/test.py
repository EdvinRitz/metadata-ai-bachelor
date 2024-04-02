#Relevant metadata:
#Namn, dokumenttyp, process (3.4 medlemstöd), organisation (vem TAM har gjort dealen med), Klassificering (Publikt, internt), 
#Personuppgifter (inga, finns, känsliga), dokumentdatum, person som gav ut.

import torch
from transformers import AutoTokenizer, AutoModelForTokenClassification
from transformers import pipeline
from shareplum import Site, Office365
from shareplum.site import Version
import docx
from lxml import etree
import zipfile
import re

site_url = ''

documentText = ''

modified = ''


tokenizer = AutoTokenizer.from_pretrained("xlm-roberta-large-finetuned-conll03-english")
model = AutoModelForTokenClassification.from_pretrained("xlm-roberta-large-finetuned-conll03-english",
    ignore_mismatched_sizes=True  # Suppress warnings about mismatched sizes
    )
classifier = pipeline("ner", model=model, tokenizer=tokenizer)

qa_model = pipeline("question-answering", "timpal0l/mdeberta-v3-base-squad2")

questionDocumentType = "Vad är titeln för dokumentet?"
questionProcess = "Vilken process beskriver dokumentet?"
questionPublisher = "Vem är ansvarig utgivare?"

authcookie = Office365('', username='', password='').GetCookies()
site = Site(site_url, version=Version.v365, authcookie=authcookie)

#for file_info in folder.files:
    #print(file_info['Name'])
    #print(file_info['ServerRelativeUrl'])

#folder.upload_file('Test', 'test.txt')

def merge_and_clean_entities(entities, text):
    if not entities:
        return []

    # Sort entities by their start position
    entities.sort(key=lambda x: x['start'])

    merged_entities = []
    current_entity = entities[0].copy()  # Make a copy to avoid mutating the original

    for next_entity in entities[1:]:
        # Extract text between the current and next entity
        gap_text = text[current_entity['end']:next_entity['start']]
        
        # Lowercase comparison to decide on merging
        gap_text_lower = gap_text.lower()
        if gap_text_lower in [' ', ''] or re.match(r'^[\s,.]+$', gap_text_lower):
            # If entities are adjacent or separated by acceptable characters, merge them
            current_entity['end'] = next_entity['end']
        else:
            # Add the entity text (preserving original casing) before starting a new entity
            merged_entities.append(text[current_entity['start']:current_entity['end']])
            current_entity = next_entity.copy()

    # Don't forget the last entity
    merged_entities.append(text[current_entity['start']:current_entity['end']])

    # Clean and deduplicate while preserving original casing
    cleaned_entities = [clean_entity(entity) for entity in merged_entities]
    unique_entities_list = list(set(cleaned_entities))

    return unique_entities_list

# Define a clean_entity function that doesn't lower case but still removes unwanted characters
def clean_entity(entity):
    # Remove leading/trailing punctuation; keep internal punctuation and original casing
    return re.sub(r'(^\W+|\W+$)', '', entity)

    
def process_docx_file(file_content, file_name):
    # Save the downloaded document
    with open(file_name, 'wb') as word_file:
        word_file.write(file_content)

    # Open and process the document
    doc = docx.Document(file_name)
    documentText = ""
    for para in doc.paragraphs:
        documentText += para.text + "\n"
    
    print('start')
    # Here you can do something with the extracted text
    print(f"Processed {file_name}")

    # Sample execution with your classifier and documentText
    resOrgs = classifier(documentText)

    # Assuming classifier results are stored in res1, and it includes 'start' and 'end' positions
    org_entities = [entity for entity in resOrgs if entity['entity'] in ('I-ORG', 'B-ORG')]

    # Preparing entities for merging, including word extraction
    for entity in org_entities:
        entity['word'] = documentText[entity['start']:entity['end']]

    # Merge, clean, and deduplicate entities
    merged_and_cleaned_org_entities = merge_and_clean_entities(org_entities, documentText)

    print(merged_and_cleaned_org_entities)

    resPublisher = qa_model(question = questionPublisher, context = documentText)
    resDocumentType = qa_model(question = questionDocumentType, context = documentText)
    resProcess = qa_model(question = questionProcess, context = documentText)

    #answer_DocumentType = resPublisher['answer']

    print('Ansvarig person är: ' + resPublisher['answer'])
    print('Dokumenttyp är: ' + resDocumentType['answer'])
    print('Processen är: ' + resProcess['answer'])

        # Extract metadata
    with zipfile.ZipFile(file_name, 'r') as docx_zip:
        # Extracting core properties including author and dates
        with docx_zip.open('docProps/core.xml') as core_xml:
            tree = etree.parse(core_xml)
            namespaces = {
            'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
            'dc': 'http://purl.org/dc/elements/1.1/',
            'dcterms': 'http://purl.org/dc/terms/',
            'dcmitype': 'http://purl.org/dc/dcmitype/',
            'xsi': 'http://www.w3.org/2001/XMLSchema-instance'
            }
        
            created = tree.find('.//dcterms:created', namespaces)
            modified = tree.find('.//dcterms:modified', namespaces)
            creator = tree.find('.//dc:creator', namespaces)
        
            print(f"Creation date: {created.text if created is not None else 'Not available'}")
            print(f"Last modified date: {modified.text if modified is not None else 'Not available'}")
            print(f"Author: {creator.text if creator is not None else 'Not available'}")

        # Extracting application (program name) properties
        with docx_zip.open('docProps/app.xml') as app_xml:
            tree = etree.parse(app_xml)
            namespaces = {
                'ep': 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties',
                'vt': 'http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes'
            }
        
            application = tree.find('.//ep:Application', namespaces)
        
            print(f"Program Name: {application.text if application is not None else 'Not available'}")

    print('end')
    

def explore_and_process_docx(folder_path, file_limit):
    # Check if the file processing limit has been reached
    if processed_files_counter['count'] >= file_limit:
        return  # Exit the function
    
    try:
        folder = site.Folder(folder_path)
        explored_folders_counter['count'] += 1
        # Process each file in the current folder
        for file_info in folder.files:
            file_name = file_info['Name']
            if file_name.endswith('.docx'): 
                print(f"Found .docx file: {file_name}")
                downloaded_doc = folder.get_file(file_name)
                process_docx_file(downloaded_doc, file_name)
                processed_files_counter['count'] += 1
        
        # Attempt to list and explore subfolders
        subfolders = folder.folders
        for subfolder_name in subfolders:
            # Construct the path for the subfolder
            subfolder_path = f"{folder_path}/{subfolder_name}" if folder_path else subfolder_name
            print(f"Exploring subfolder: {subfolder_path}")
            explore_and_process_docx(subfolder_path, file_limit)
    except Exception as e:
        print(f"Error exploring {folder_path}: {e}")

# Start exploration from the root folder path
root_folder_path = 'Delade dokument'
processed_files_counter = {'count': 0}
explored_folders_counter = {'count': 0}
file_limit = 10
explore_and_process_docx(root_folder_path, file_limit)

print('Files processed: ' + str(processed_files_counter['count']))
print('Folders explored: ' + str(explored_folders_counter['count']))

#explore_and_process_docx(folder)

#print(download)

#with open(word_file_name, 'wb') as word_file:
    #word_file.write(downloaded_doc)

#doc = docx.Document(word_file_name)

""" for para in doc.paragraphs:
    documentText += para.text + "\n"
    #print(para.text)

#print(documentText)
# Extract metadata
with zipfile.ZipFile(word_file_name, 'r') as docx_zip:
    # Extracting core properties including author and dates
    with docx_zip.open('docProps/core.xml') as core_xml:
        tree = etree.parse(core_xml)
        namespaces = {
        'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
        'dc': 'http://purl.org/dc/elements/1.1/',
        'dcterms': 'http://purl.org/dc/terms/',
        'dcmitype': 'http://purl.org/dc/dcmitype/',
        'xsi': 'http://www.w3.org/2001/XMLSchema-instance'
        }
        
        created = tree.find('.//dcterms:created', namespaces)
        modified = tree.find('.//dcterms:modified', namespaces)
        creator = tree.find('.//dc:creator', namespaces)
        
        print(f"Creation date: {created.text if created is not None else 'Not available'}")
        print(f"Last modified date: {modified.text if modified is not None else 'Not available'}")
        print(f"Author: {creator.text if creator is not None else 'Not available'}")

    # Extracting application (program name) properties
    with docx_zip.open('docProps/app.xml') as app_xml:
        tree = etree.parse(app_xml)
        namespaces = {
            'ep': 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties',
            'vt': 'http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes'
        }
        
        application = tree.find('.//ep:Application', namespaces)
        
        print(f"Program Name: {application.text if application is not None else 'Not available'}")
 """



#"Vad är detta för dokument?" "Vad är detta för process?" "Vilket företag gäller dokumentet?" "Vad är dokumentdatumet?" "Vem är referensperson?" 



#context = documentText

#resPublisher = qa_model(question = questionPublisher, context = context)
#resDocumentType = qa_model(question = questionDocumentType, context = context)
#resProcess = qa_model(question = questionProcess, context = context)

#answer_DocumentType = resPublisher['answer']

#print(resPublisher)
#print(resDocumentType)
#print(resProcess)
#print(answer_DocumentType)
